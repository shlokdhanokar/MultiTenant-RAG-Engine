"""
DOCX layout parser.

Word documents carry real structural metadata — paragraph style names like
"Heading 1"/"Heading 2" — so heading detection here is exact rather than the
font-size heuristic the PDF parser is forced to use. Output is normalized to
the same shape as analyze_document_layout() so the downstream chunker,
image mapper, and embedding stages are format-agnostic.
"""
import logging

logger = logging.getLogger(__name__)


MAX_HEADING_WORDS = 12


def _paragraph_font_size(paragraph):
    """Dominant explicit font size in points, or None if the paragraph inherits it."""
    sizes = [r.font.size.pt for r in paragraph.runs if r.font is not None and r.font.size is not None]
    if not sizes:
        return None
    return max(set(sizes), key=sizes.count)


def _is_bold(paragraph):
    """True when every run carrying text is bold."""
    runs = [r for r in paragraph.runs if (r.text or "").strip()]
    return bool(runs) and all(r.bold for r in runs)


def _body_font_size(paragraphs):
    """
    Median explicit font size across the document, used as the body-text
    baseline. Mirrors the PDF parser's approach so both formats classify
    headings on the same principle.
    """
    sizes = sorted(
        s for s in (_paragraph_font_size(p) for p in paragraphs if (p.text or "").strip())
        if s is not None
    )
    if not sizes:
        return None
    mid = len(sizes) // 2
    return (sizes[mid - 1] + sizes[mid]) / 2 if len(sizes) % 2 == 0 else sizes[mid]


def _classify_paragraph(paragraph, body_size=None):
    """
    Map a Word paragraph to the pipeline's block types.

    Semantic styles ("Heading 1", "Title") are authoritative when present. Many
    real-world documents skip them and format headings directly instead, so a
    paragraph that is bold, short, and visibly larger than body text is also
    treated as a heading — otherwise those documents ingest as one flat blob
    with no topic structure at all.
    """
    style_name = (paragraph.style.name if paragraph.style is not None else "") or ""
    normalized = style_name.strip().lower()

    if normalized in ("title", "heading 1"):
        return "Topic"
    if normalized.startswith("heading"):
        # Heading 2..9 all act as subtopics; the chunker only distinguishes two levels.
        return "Subtopic"

    # Fallback for directly-formatted headings.
    text = (paragraph.text or "").strip()
    size = _paragraph_font_size(paragraph)
    if (
        body_size
        and size
        and len(text.split()) <= MAX_HEADING_WORDS
        and _is_bold(paragraph)
    ):
        if size >= body_size * 1.40:
            return "Topic"
        if size >= body_size * 1.15:
            return "Subtopic"

    return "Paragraph"


def analyze_docx_layout(docx_bytes: bytes):
    """
    Extracts structured blocks and embedded images from a .docx file.

    Word has no page concept available without rendering, so page_number is
    reported as 0 throughout; the chunker only uses it for ordering, and
    document order is already preserved here.
    """
    import io
    from docx import Document

    try:
        document = Document(io.BytesIO(docx_bytes))
    except Exception as e:
        logger.error(f"failed to load docx from memory: {e}")
        return None

    body_size = _body_font_size(document.paragraphs)

    semantic_blocks = []
    for index, paragraph in enumerate(document.paragraphs):
        text = (paragraph.text or "").strip()
        if not text:
            continue

        semantic_blocks.append({
            "page_number": 0,
            "text": text,
            "font_size": _paragraph_font_size(paragraph) or 0,
            "bbox": (0, index, 0, index),  # preserves document order for image anchoring
            "type": _classify_paragraph(paragraph, body_size),
        })

    # Tables carry content that would otherwise be dropped entirely.
    for table in document.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                semantic_blocks.append({
                    "page_number": 0,
                    "text": row_text,
                    "font_size": 0,
                    "bbox": (0, len(semantic_blocks), 0, len(semantic_blocks)),
                    "type": "Paragraph",
                })

    images = []
    for rel in document.part.rels.values():
        if "image" not in rel.reltype:
            continue
        try:
            blob = rel.target_part.blob
        except Exception:
            continue
        ext = rel.target_part.partname.ext.lstrip(".") if hasattr(rel.target_part, "partname") else "png"
        images.append({
            "page_number": 0,
            "image_bytes": blob,
            "image_ext": ext,
            "width": 0,
            "height": 0,
            # Word doesn't expose reliable layout coordinates without rendering,
            # so images anchor to the end of the document rather than guessing.
            "bbox": (0, len(semantic_blocks), 0, len(semantic_blocks)),
        })

    return {
        "total_pages": 1,
        "semantic_blocks": semantic_blocks,
        "images": images,
    }
